# -*- coding: utf-8 -*-
"""
    Security Knowledge Framework is an expert system application
    that uses OWASP Application Security Verification Standard, code examples,
    helps developers in pre-development and post-development.
    Copyright (C) 2015  Glenn ten Cate, Riccardo ten Cate

    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU Affero General Public License as
    published by the Free Software Foundation, either version 3 of the
    License, or (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU Affero General Public License for more details.

    You should have received a copy of the GNU Affero General Public License
    along with this program. If not, see <http://www.gnu.org/licenses/>.
"""

import contextlib, traceback
import os, markdown, datetime, string, base64, re, sys, re, requests, mimetypes, smtplib
from OpenSSL import SSL, rand
from docx import Document
from BeautifulSoup import BeautifulSoup
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from functools import wraps
from sqlite3 import dbapi2 as sqlite3
from flask.ext.bcrypt import Bcrypt
from flask import Flask, request, session, g, redirect, url_for, abort, \
     render_template, flash, Markup, make_response



# create the application
app = Flask(__name__)

"""Set up bcrypt for passwords encrypting"""
bcrypt = Bcrypt(app)


def add_response_headers(headers={}):
    """This decorator adds the headers passed in to the response"""
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            resp = make_response(f(*args, **kwargs))
            h = resp.headers
            for header, value in headers.items():
                h[header] = value
            return resp
        return decorated_function
    return decorator


def security(f):
    """This decorator passes multiple security headers and checks log file to block users"""
    # return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache','Strict-Transport-Security': 'max-age=16070400; includeSubDomains', 'Server': 'Security Knowledge Framework'})(f)
    return add_response_headers({'X-Frame-Options': 'deny', 'X-XSS-Protection': '1', 'X-Content-Type-Options': 'nosniff', 'Cache-Control': 'no-store, no-cache', 'Server': 'Security Knowledge Framework'})(f)


def check_token():
    """Checks the submitted CSRF token"""
    if not session.get('csrf_token') == request.form['csrf_token']:
        log("User supplied CSRF token not valid", "FAIL", "HIGH")
        session.clear()
        return abort(500)(f)


def generate_pass():
    chars = string.letters + string.digits + '+/'
    assert 256 % len(chars) == 0  # non-biased later modulo
    PWD_LEN = 12
    generated_pass = ''.join(chars[ord(c) % len(chars)] for c in os.urandom(PWD_LEN))
    return generated_pass


def random_token(tokenBytes):
    #Create random token
    rand.cleanup()
    Random_token_raw = rand.bytes(int(tokenBytes))
    Random_token = base64.b64encode(Random_token_raw)
    result = re.sub("==", "", Random_token)
    return result


def log(message, value, threat):
    """Create log file and write events triggerd by the user
    The variables: message can be everything, value contains FAIL or SUCCESS and threat LOW MEDIUM HIGH"""
    now = datetime.datetime.now()
    dateLog = now.strftime("%Y-%m")
    dateTime = now.strftime("%Y-%m-%d %H:%M")
    ip = request.remote_addr
    fullmessage = dateTime +' '+ message +' ' + ' ' + value + ' ' + threat + ' ' + ip
    sys.stderr.write(fullmessage)
    try:
        file = open('logs/'+dateLog+'.txt', 'a+')
    except IOError:
        # If not exists, create the file
        file = open('logs/'+dateLog+'.txt', 'w+')
    file.write(fullmessage + "\r\n")
    file.close()


def blockUsers():
    """Check the log file and based on the FAIL items block a user"""
    dateLog  = datetime.datetime.now().strftime("%Y-%m")
    count = 0
    try:
        read = open(os.path.join(app.root_path, 'logs/'+dateLog+'.txt'), 'a+')
    except IOError:
        # If not exists, create the file
        read = open(os.path.join(app.root_path, 'logs/'+dateLog+'.txt'), 'w+')
    for line in read:
        match = re.search('FAIL', line)
        # If-statement after search() tests if it succeeded
        if match:
            count += 1
            str(count)
            if count > 11:
                sys.exit('Due to to many FAILED logs in your logging file we have the suspicion your application has been under attack by hackers. Please check your log files to validate and take repercussions. After validation clear your log or simply change the FAIL items to another value.')


def whiteList(range, value, countLevel):
    match = re.findall(range, value)
    if match:
        return True
    else:
        log("User supplied value not in the range " + range, "FAIL", "MEDIUM")
        countAttempts(countLevel)
        abort(406)
        return False


def valAlphaNum(value, countLevel):
    return whiteList(r'^([ a-zA-Z0-9_.-]*)$', value, countLevel)


def valNum(value, countLevel):
    return whiteList(r'^([0-9]+)$', value, countLevel)


def valBool(value, countLevel):
    return whiteList(r'^(true|false)$', value, countLevel)


#secret key for flask internal session use
rand.cleanup()
secret_key = rand.bytes(512)

mimetypes.add_type('image/svg+xml', '.svg')
bindaddr = '127.0.0.1';

# Load default config and override config from an environment variable
# You can also replace password with static password:  PASSWORD='pass!@#example'
app.config.update(dict(
    DATABASE=os.path.join(app.root_path, 'skf.db'),
    DEBUG=False,
    SECRET_KEY=secret_key,
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY = True
))


@app.context_processor
def inject_year():
    return dict(year=datetime.datetime.now().strftime("%Y"))


def connect_db():
    """Connects to the specific database."""
    rv = sqlite3.connect(app.config['DATABASE'])
    rv.row_factory = sqlite3.Row
    return rv


def init_db():
    """Initializes the database."""
    with app.open_resource('schema.sql') as f:
        with contextlib.closing(get_db()) as con:
            con.cursor().executescript(f.read())


@app.cli.command('initdb')
def initdb_command():
    """Creates the database tables."""
    init_db()
    print('Initialized the database.')


def get_db():
    """Opens a new database connection if there is none yet for the
    current application context.
    """
    if False:
        if not hasattr(g, 'sqlite_db'):
            g.sqlite_db = connect_db()
        return g.sqlite_db
    else:
        return connect_db()


@app.teardown_appcontext
def close_db(error):
    """Closes the database again at the end of the request."""
    # The database will have been closed by the end of the request, thanks to
    # contextlib.closing().  Besides, the database appears closed at the
    # beginning of the request, probably due to this teardown preserving 
    # the closed connection as g.sqlite_db.
    if False:
        if hasattr(g, 'sqlite_db'):
            g.sqlite_db.close()
            delattr(g, 'sqlite_db')


@app.errorhandler(400)
def handle_error(exc):
    log(('Bad request: %s\n%s' % (exc, traceback.format_exc())).encode("string_escape"), "SUCCESS", "LOW")
    app.logger.exception('Bad request')
    return "Bad request"


@app.errorhandler(Exception)
def handle_error(exc):
    log(('Unhandled exception: %s\n%s' % (exc, traceback.format_exc())).encode("string_escape"), "SUCCESS", "LOW")
    app.logger.exception('Unhandled exception')
    return "Internal server error"


def get_filepaths(directory):
    """
    This function will generate the file names in a directory
    tree by walking the tree either top-down or bottom-up. For each
    directory in the tree rooted at directory top (including top itself),
    it yields a 3-tuple (dirpath, dirnames, filenames).
    """
    file_paths = []
    for root, directories, files in os.walk(directory):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_paths.append(filepath)
    return file_paths


def get_num(x):
    """get numbers from a string"""
    return int(''.join(ele for ele in x if ele.isdigit()))


def check_version():
    try:
        r = requests.get("http://raw.githubusercontent.com/blabla1337/skf-flask/master/setup.py")
        items_remote = r.content.split(",")
        version_remote = items_remote[1]
        version_remote = version_remote.replace(version_remote[:14], '')
        version_remote = version_remote[:-1]
        with open ("version.txt", "r") as myfile:
            version_local = myfile.read().replace('\n', '')

        if version_local == version_remote:
            return True
        else:
            return False
    except:
        return False


def get_version():
    with open ("version.txt", "r") as myfile:
        version_final = myfile.read().replace('\n', '')
    return version_final


def projects_functions_techlist():
    """get list of technology used for creating project functions"""
    if not session.get('logged_in'):
        abort(401)
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT techID, techName, vulnID from techhacks ORDER BY techID DESC').fetchall()
    return entries


@app.route('/')
@security
def show_landing():
    """show the loging page and set default code language"""
    rand.cleanup()
    csrf_token_raw = rand.bytes(128)
    csrf_token = base64.b64encode(csrf_token_raw)
    session['csrf_token'] = csrf_token
    session['code_lang'] = "php"

    return render_template('login.html', csrf_token=session['csrf_token'])


@app.route('/dashboard', methods=['GET'])
@security
def dashboard():
    """show the landing page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /dashboard", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    version_check = check_version()
    version = get_version()
    return render_template('dashboard.html', version=version, version_check=version_check)


@app.route('/first-login', methods=['GET'])
@security
def first_login():
    version_check = check_version()
    version = get_version()
    return render_template('first-login.html', version=version, version_check=version_check)


"""create account for a user"""
@app.route('/create-account', methods=['GET', 'POST'])
@security
def create_account():
    """validate the login data for access dashboard page"""
    error = None
    with contextlib.closing(get_db()) as con:
        if request.method == 'POST':
            """Username, password, token, email from form"""
            token  = request.form['token']
            email  = request.form['email']
            password  = request.form['password']

            #hash the password with Bcrypt, does autosalt
            hashed = bcrypt.generate_password_hash(password,14)

            #Do DB query also check for access
            check = con.execute('SELECT accessToken, userID, activated from users where email=? AND accessToken=?',
                                [email, token]).fetchall()
            for verify in check:
                userID = verify[1]
                if verify[2] == "false":
                    if str(verify[0]) == token:
                        #update the counter and blocker table with new values
                        with con as cur:
                            cur.execute('UPDATE users SET access=?, password=?, activated=? WHERE accessToken=? AND userID=?',
                                           ["true", hashed, "true", token , userID])
                        #Insert record in counter table for the counting of malicious inputs
                        with con as cur:
                            cur.execute('INSERT INTO counter (userID, countEvil, block) VALUES (?, ?, ?)',
                                                [userID, 0, 0])

                        #Create standard group  for this user to assign himself to
                        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        with con as cur:
                            cur.execute('INSERT INTO groups (ownerID, groupName, timestamp) VALUES (?, ?, ?)',
                                                [userID, "privateGroup", date])

                        #Select this groupID so we can assign the user to this group automatically
                        group = con.execute('SELECT groupID from groups where ownerID=?',
                                                [userID]).fetchall()
                        for theID in group:
                            groupID = theID[0]

                        #Now we assign the user to the group
                        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                        with con as cur:
                            cur.execute('INSERT INTO groupMembers (userID, groupID, ownerID) VALUES (?, ?, ?)',
                                                [userID, groupID, userID])

            if not check:
                #if not the right pin, the user account wil be deleted if not exsisting
                with con as cur:
                    con.execute('DElETE FROM users where email=? AND activated=?',
                            [email, "false"])

        return render_template('login.html', error=error)


"""First comes the method for login"""
@app.route('/login', methods=['GET', 'POST'])
@security
def login():
    """validate the login data for access dashboard page"""
    error = None
    with contextlib.closing(get_db()) as con:
        if request.method == 'POST':
            """Username and password from form"""
            username = request.form['username']
            password = request.form['password']

            #Do DB query also check for access
            check = con.execute('SELECT access from users where userName=?',
                                [username]).fetchall()
            for verify in check:
                if verify[0] == "false":
                    return render_template('warning.html', error=error)

            #Do DB query also check for access
            entries = con.execute('SELECT u.userID, u.privilegeID, u.userName, u.password, u.access, priv.privilegeID, priv.privilege from users as u JOIN privileges AS priv ON priv.privilegeID = u.privilegeID where username=? AND access="true"',
                                [username]).fetchall()
            for entry in entries:
                passwordHash = entry[3]
                userID = entry[0]
                #Do encryption
                if bcrypt.check_password_hash(passwordHash, password):
                    log("Valid username/password submit", "SUCCESS", "HIGH")
                    rand.cleanup()
                    csrf_token_raw = rand.bytes(128)
                    csrf_token = base64.b64encode(csrf_token_raw)
                    session['logged_in'] = True
                    session['userID'] = userID
                    session['csrf_token'] = csrf_token
                    session['code_lang'] = "php"
                    session['userName'] = entry[2]
                    valAlphaNum(session['userName'], 12)
                    session['permissions'] = entry[6]
                    version_check = check_version()
                    version = get_version()

                    #Do DB query also check for access
                    groupID = con.execute('SELECT groupID from groups WHERE groupName=? AND ownerID=?',
                                ["privateGroup", session['userID']]).fetchall()
                    for entry in groupID:
                        session['privateGroup'] = entry[0]
                    return render_template('dashboard.html', version=version, version_check=version_check)
                else:
                    log("invalid login submit", "FAIL", "HIGH")
    return render_template('login.html', error=error)


def countAttempts(counter):
    """We count hacking attempts and block the user if structural"""
    if not session.get('logged_in'):
        abort(401)

    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT * FROM counter where userID=?',
                            [session['userID']]).fetchall()
        counterDB = 0
        blockDB = 0
        for entry in entries:
            counterDB = entry[2]
            blockDB = entry[3]

        updateCount = counterDB + counter
        updateBlock = blockDB + counter
        redirect = False

        if updateCount >= 3:
            countUpdate = 0
            redirect = True

        if updateBlock >= 12:
            redirect = True
            with con as cur:
                cur.execute('UPDATE users SET access=? WHERE userID=?',
                   ["false", session['userID']])
            renderwhat = "/warning.html"

        #update the counter and blocker table with new values
        with con as cur:
            cur.execute('UPDATE counter SET countEvil=?, block=? WHERE userID=?',
                [updateCount, updateBlock, session['userID']])

        if redirect:
            log( "Authenticated session destroyed by counter class", "SUCCESS", "LOW")
            # TO-DO turn on again
            #session.pop('logged_in', None)
            #session.clear()


"""Here is the method for the database enforced privilege based authentication"""
def permissions(fromFunction):
    with contextlib.closing(get_db()) as con:
        """Do DB query to see if username exists"""
        entries = con.execute('SELECT a.username, a.userID, a.password, a.privilegeID, b.privilegeID, b.privilege FROM users as a JOIN privileges as b ON a.privilegeID = b.privilegeID WHERE a.userID =? and a.access="true" ',
                                           [session['userID']]).fetchall()
        perms = ''
        for entry in entries:
            perms = entry[5]

        permissionsGranted = string.split(perms, ':')
        permissionsNeeded  = string.split(fromFunction, ':')

        count = len(permissionsNeeded)
        counthits = 0

        for val in permissionsGranted:
                if val in fromFunction:
                    counthits +=1
        if counthits >= count:
            return perms
        else:
            log( "User tries to reach functions out of bound no restrictions!!", "FAIL", "HIGH")
            abort(401)


@app.route('/logout', methods=['GET', 'POST'])
@security
def logout():
    """logout and destroy session"""
    log( "Authenticated session destroyed", "SUCCESS", "LOW")
    session.pop('logged_in', None)
    session.clear()
    return redirect("/")


@app.route('/code/<code_lang>', methods=['GET'])
@security
def set_code_lang(code_lang):
    """set a code language: php java python perl"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /code", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    allowed = "php java asp"
    valAlphaNum(code_lang, 12)
    found = allowed.find(code_lang)
    if found != -1:
        session['code_lang'] = code_lang
    return redirect(url_for('code_examples'))


@app.route('/code-examples', methods=['GET'])
@security
def code_examples():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /code-examples", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    items = []
    id_items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            basepath = os.path.basename(path)
            baseelems = basepath.split("-")
            id_item = get_num(baseelems[0])
            kb_name_uri = baseelems[-3]
            kb_name = kb_name_uri.replace("_", " ")
            items.append(kb_name)
            id_items.append(id_item)
    return render_template('code-examples.html', items=items, id_items=id_items)


@app.route('/code-item', methods=['POST'])
@security
def show_code_item():
    """show the coding examples page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /code-item", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    valNum(request.form['id'], 12)
    id = int(request.form['id'])
    items = []
    full_file_paths = []
    allowed = set(string.ascii_lowercase + string.ascii_uppercase + '.')
    if set(session['code_lang']) <= allowed:
        full_file_paths = get_filepaths(os.path.join(app.root_path, "markdown/code_examples/"+session['code_lang']))
        for path in full_file_paths:
            basepath = os.path.basename(path)
            if id == get_num(basepath.split("-")[0]):
                with open(path, 'r') as codef:
                    codemd = codef.read()
                content = Markup(markdown.markdown(codemd))
    return render_template('code-examples-item.html', **locals())


@app.route('/kb-item', methods=['POST'])
@security
def show_kb_item():
    """show the knowledge base search result page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /kb-item", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    valNum(request.form['id'], 12)
    id = int(request.form['id'])
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        basepath = os.path.basename(kbpath)
        if id == get_num(basepath.split("-")[0]):
            with open(kbpath, 'r') as kbpathf:
                kbmd = kbpathf.read()
            content = Markup(markdown.markdown(kbmd))
    return render_template('knowledge-base-item.html', **locals())


@app.route('/knowledge-base-api', methods=['GET'])
@security
def show_kb_api():
    """show the knowledge base items page"""
    log( "User access page /knowledge-base-api", "SUCCESS", "HIGH")
    full_file_paths = []
    content = []
    kb_name = []
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        filetmp = open(kbpath, 'r').read()
        filetmp2 = filetmp.replace("-------", "")
        filetmp3 = filetmp2.replace("**", "")
        filetmp4 = filetmp3.replace("\"", "")
        filetmp5 = filetmp4.replace("\t", "")
        content.append(filetmp5.replace("\n", " "))
        basepath = os.path.basename(kbpath)
        kb_name_uri = basepath.split("-")[-3]
        kb_name.append(kb_name_uri.replace("_", " "))
    return render_template('knowledge-base-api.html', **locals())


@app.route('/knowledge-base', methods=['GET'])
@security
def knowledge_base():
    """Shows the knowledge base markdown files."""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /knowledge-base", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    items = []
    id_items = []
    kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
    for kbpath in kb_paths:
        basepath = os.path.basename(kbpath)
        elems = basepath.split("-")
        id_item = get_num(elems[0])
        kb_name_uri = elems[-3]
        kb_name = kb_name_uri.replace("_", " ")
        items.append(kb_name)
        id_items.append(id_item)
    return render_template('knowledge-base.html', items=items, id_items=id_items)


@app.route('/users-new', methods=['GET'])
@security
def user_new():
    """show the create new project page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /user-new", "FAIL", "HIGH")
        abort(401)
    permissions("manage")
    return render_template('users-new.html', csrf_token=session['csrf_token'])


@app.route('/users-add', methods=['POST'])
@security
def users_add():
    """add a new project to database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /users-add", "FAIL", "HIGH")
        abort(401)
    permissions("manage")
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            userName = request.form['username']
            email    = request.form['email']
            privID   = request.form['privID']
            pincode  = request.form['pincode']
            valAlphaNum(username, 1)
            valNum(privID, 12)
            valNum(pincode, 12)

            cur.execute('INSERT INTO users (privilegeID, userName, email, password, access, accessToken, activated) VALUES (?, ?, ?, ?, ?, ?, ?)',
                       [privID, userName, email, "none", "false", pincode, "false"])

    return redirect(url_for('users_manage'))


@app.route('/users-manage', methods=['GET'])
@security
def users_manage():
    """show the project list page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /group-manage", "FAIL", "HIGH")
        abort(401)
    permissions("manage")
    with contextlib.closing(get_db()) as con:
        users = con.execute('SELECT u.userID, u.userName, u.email, u.privilegeID, u.access, p.privilegeID, p.privilege from users as u JOIN privileges as p ON p.privilegeID = u.privilegeID').fetchall()

    return render_template('users-manage.html', users=users, csrf_token=session['csrf_token'])


@app.route('/user-access', methods=['POST'])
@security
def user_access():
    """add a new project to database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /assign-group", "FAIL", "HIGH")
        abort(401)
    permissions("manage")
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            access = request.form['access']
            valBool(access, 12)
            userID = request.form['userID']
            valNum(userID, 12)
            cur.execute('UPDATE users SET access=? WHERE userID=?',
                           [access, userID])
            cur.execute('UPDATE counter SET countEvil=? AND block=? WHERE userID=?',
                           [0, 0, userID])

    return redirect(url_for('users_manage'))


@app.route('/group-new', methods=['GET'])
@security
def group_new():
    """show the create new project page"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /group-new", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    return render_template('group-new.html', csrf_token=session['csrf_token'])

@app.route('/group-add', methods=['POST'])
@security
def group_add():
    """add a new project to database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /group-add", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        inputName = request.form['groupName']
        valAlphaNum(groupName, 3)
        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        with con as cur:
            cur.execute('INSERT INTO groups (timestamp, groupName, ownerID) VALUES (?, ?, ?)',
                   [date, inputName, session['userID']])
        # select the latest group in order to check id
        group = con.execute('SELECT groupID from groups WHERE timestamp=? AND ownerID=?',
                            [date, session['userID']]).fetchall()
        groupID = None
        for value in group:
            groupID = value[0]

        # insert this back into groupMembers table so the user is added to group
        with con as cur:
            cur.execute('INSERT INTO groupMembers (userID, groupID, ownerID, timestamp) VALUES (?, ?, ?, ?)',
                   [session['userID'], groupID, session['userID'], date])
    return redirect(url_for('group_manage'))


@app.route('/group-users', methods=['GET'])
@security
def group_users():
    """show the project list page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /group-users", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    with contextlib.closing(get_db()) as con:
        groups = con.execute('SELECT * from groups where ownerID=? and groupName !=? ',
                              [session['userID'], "privateGroup"]).fetchall()

        """Select all users for adding to group"""
        users = con.execute('SELECT username, userID from users').fetchall()

        """select users by assigned groups for display"""
        summary = con.execute('SELECT u.username, u.userID, g.groupName, g.groupID, m.groupID, m.userID, m.timestamp, g.ownerID from users as u JOIN groups AS g ON g.groupID = m.groupID JOIN groupMembers as m ON u.userID = m.userID  WHERE g.ownerID=? AND u.userName !=? ORDER BY g.groupName ',
                                       [session['userID'], session['userName']]).fetchall()

    return render_template('group-users.html', groups=groups, users=users, summary=summary, csrf_token=session['csrf_token'])


@app.route('/group-add-users', methods=['POST'])
@security
def group_add_users():
    """add a project function"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-function-add", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    check_token()
    valNum(request.form['groupName'], 12)
    groupID = request.form['groupName']

    """Check is submitted groupID is owned by user"""
    with contextlib.closing(get_db()) as con:
        owner = con.execute('SELECT groupID from groups where ownerID=?',
                                       [session['userID']]).fetchall()
        for val in owner:
            if int(groupID) == int(val[0]):
                f = request.form
                for key in f.keys():
                    for value in f.getlist(key):
                        found = key.find("test")
                        if found != -1:
                            date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                            items = value.split("-")
                            userID = items[0]
                            valNum(userID, 12)
                            with con as cur:
                                cur.execute('INSERT INTO groupMembers (timestamp, groupID, userID, ownerID) VALUES (?, ?, ?, ?)',
                                       [date, groupID, userID, session['userID']])
    redirect_url = '/group-users'
    return redirect(redirect_url)


@app.route('/user-del', methods=['POST'])
@security
def user_del():
    """delete project from database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /user-del", "FAIL", "HIGH")
        abort(401)
    permissions("delete")
    check_token()
    userID  = request.form['userID']
    valNum(userID, 12)

    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM users WHERE userID=?",
                [userID])
    return redirect("/users-manage")


@app.route('/group-manage', methods=['GET'])
@security
def group_manage():
    """show the project list page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /group-manage", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    with contextlib.closing(get_db()) as con:
        groups = con.execute('SELECT * from groups where ownerID=? and groupName !=? ',
                          [session['userID'], "privateGroup"]).fetchall()

    return render_template('group-manage.html', groups=groups, csrf_token=session['csrf_token'])


@app.route('/group-del', methods=['POST'])
@security
def group_del():
    """delete project from database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /group-del", "FAIL", "HIGH")
        abort(401)
    permissions("manage")
    check_token()
    groupID = request.form['groupID']
    valNum(groupID, 12)

    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM groups WHERE groupID=? AND ownerID=?",
                [groupID, session['userID']])
    return redirect("/group-manage")


@app.route('/project-new', methods=['GET'])
@security
def projects():
    """show the create new project page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-new", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    return render_template('project-new.html', csrf_token=session['csrf_token'])


@app.route('/project-add', methods=['POST'])
@security
def add_entry():
    """add a new project to database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-add", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        inputName = request.form['inputName']
        inputVersion = request.form['inputVersion']
        inputDesc = request.form['inputDesc']
        # valAlphaNum(inputName, 1)
        # valNum(inputVersion, 1)
        date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        with con as cur:
            cur.execute('INSERT INTO projects (timestamp, projectName, projectVersion, projectDesc, userID, ownerID, groupID) VALUES (?, ?, ?, ?, ?, ?, ?)',
                   [date, inputName, inputVersion, inputDesc, session['userID'],  session['userID'], session['privateGroup']])
    return redirect(url_for('project_list'))


@app.route('/assign-group', methods=['POST'])
@security
def assign_group():
    """add a new project to database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /assign-group", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    check_token()
    with contextlib.closing(get_db()) as con:
        groupID   = request.form['groupID']
        projectID = request.form['projectID']
        valNum(projectID, 12)
        valNum(groupID, 12)
        """Check is submitted groupID is owned by user"""
        owner = con.execute('SELECT groupID from groups where ownerID=?',
                                       [session['userID']]).fetchall()
        for val in owner:
            if int(groupID) == int(val[0]):
                with con as cur:
                    cur.execute('UPDATE projects SET groupID=? WHERE projectID=? AND userID=?',
                       [groupID, projectID, session['userID']])
    return redirect(url_for('project_list'))


@app.route('/project-del', methods=['POST'])
@security
def project_del():
    """delete project from database"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-del", "FAIL", "HIGH")
        abort(401)
    permissions("delete")
    id = request.form['projectID']
    valNum(id, 12)
    check_token()
    with contextlib.closing(get_db()) as con:
        with con as cur:
            cur.execute("DELETE FROM projects WHERE projectID=? AND userID=? AND ownerID=?",
               [id, session['userID'], session['userID']])
    return redirect("/project-list")


@app.route('/project-list', methods=['GET'])
@security
def project_list():
    """show the project list page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-list", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    with contextlib.closing(get_db()) as con:
        #First query is for the users own owned projects
        entries = con.execute('SELECT p.projectName, p.projectVersion, p.projectDESC, p.projectID, p.timestamp, p.groupID, g.groupName, g.groupID FROM projects as p JOIN groups as g ON g.groupID = p.groupID where p.userID=? ORDER BY projectID DESC',
                              [session['userID']]).fetchall()
        #select the groups which can be selected by this user
        groups = con.execute('SELECT * FROM groups WHERE ownerID=?',
                              [session['userID']]).fetchall()
    return render_template('project-list.html', entries=entries, groups=groups, csrf_token=session['csrf_token'])


@app.route('/project-shared', methods=['GET'])
@security
def project_shared():
    """show the project list page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-list", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    with contextlib.closing(get_db()) as con:
        #Here we see what projects this users was assigned to
        entries = con.execute('SELECT p.projectName, p.projectVersion, p.projectDESC, p.projectID, p.timestamp, p.groupID, p.ownerID, m.userID, m.groupID, u.userID, u.userName FROM projects as p JOIN groupMembers as m ON m.groupID = p.groupID JOIN users as u ON u.userID=p.ownerID where m.userID=? AND u.userName !=? ORDER BY p.projectID DESC',
                              [session['userID'], session['userName']]).fetchall()

    return render_template('project-shared.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/project-options/<project_id>', methods=['GET'])
@security
def projects_options(project_id):
    """show the project options landing page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-options", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    valNum(project_id, 12)
    return render_template('project-options.html', project_id=project_id, csrf_token=session['csrf_token'])


@app.route('/project-functions/<project_id>', methods=['GET'])
@security
def project_functions(project_id):
    """show the pproject functions page"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-functions", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    techlist = projects_functions_techlist()
    valNum(project_id, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT p.paramID, p.functionName, p.functionDesc, p.projectID, p.userID, p.tech, p.techVuln, p.entryDate, t.techName, proj.projectID, proj.groupID, m.userID, m.groupID FROM parameters AS p JOIN techhacks AS t ON p.tech = t.techID JOIN projects as proj ON proj.projectID = p.projectID JOIN groupMembers as m ON m.groupID = proj.groupID WHERE proj.projectID=? AND m.userID=? GROUP BY t.techName',
                      [project_id, session['userID']]).fetchall()
    return render_template('project-functions.html', project_id=project_id, techlist=projects_functions_techlist(), entries=entries, csrf_token=session['csrf_token'])


@app.route('/project-function-del', methods=['POST'])
@security
def function_del():
    """delete a project function"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /project-function-del", "FAIL", "HIGH")
        abort(401)
    permissions("delete")
    check_token()
    id = request.form['projectID']
    id_param = int(request.form['paramID'])
    valNum(id, 12)
    valNum(id_param, 12)
    with contextlib.closing(get_db()) as con:
        #First check if the user is allowed to delete this parameter
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
                                       [session['userID']]).fetchall()
        for project in projects:
            if int(id) == int(project[0]):
                with con as cur:
                    cur.execute("DELETE FROM parameters WHERE projectID=? AND paramID=?",
                                   [id, id_param])
                break
    redirect_url = "/project-functions/"+str(id)
    return redirect(redirect_url)


@app.route('/project-function-add', methods=['POST'])
@security
def add_function():
    """add a project function"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-function-add", "FAIL", "HIGH")
        abort(401)
    permissions("edit")
    check_token()
    id = request.form['project_id']
    valNum(id, 12)
    fName = request.form['functionName']
    # valAlphaNum(fName, 1)
    fDesc = request.form['functionDesc']
    # valAlphaNum(fDesc, 1)

    #Check is submitted projectID is owned by user
    with contextlib.closing(get_db()) as con:
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
                                       [session['userID']]).fetchall()
        for project in projects:
            if int(id) == int(project[0]):
                f = request.form
                for key in f.keys():
                    for value in f.getlist(key):
                        found = key.find("test")
                        if found != -1:
                            date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                            items = value.split("-")
                            techID = items[2]
                            vulnID = items[0]
                            valAlphaNum(techID, 12)
                            valAlphaNum(vulnID, 12)
                            with con as cur:
                                cur.execute('INSERT INTO parameters (entryDate, functionName, functionDesc, techVuln, tech, projectID, userID) VALUES (?, ?, ?, ?, ?, ?, ?)',
                                       [date, fName, fDesc, vulnID, techID, id, session['userID']])
    redirect_url = '/project-functions/'+str(id)
    return redirect(redirect_url)


CKELEM_PREFIX = "vulnID"
CKELEM_PREFIX_LEN = len(CKELEM_PREFIX)

@app.route('/project-checklist-add', methods=['POST'])
@security
def add_checklist():
    """add project checklist"""
    if not session.get('logged_in'):
        log("User with no valid session tries access to page /project-checklist-add", "FAIL", "HIGH")
        abort(401)
    permissions("edit")

    # https://github.com/pallets/werkzeug/issues/1068
    request.get_data()
    
    # log("Type of request: %s" % (type(request),), "SUCCESS", "LOW")
    # log("dir(request): %s" % (dir(request),), "SUCCESS", "LOW")
    # log("Type of form: %s" % (type(request.form),), "SUCCESS", "LOW")
    # for key in request.form:
    #     log("Form element: \"%s\"=\"%s\"" % (key, request.form[key],), "SUCCESS", "LOW")
    # log("Request form csrf token: %s" % (request.form["csrf_token"],), "SUCCESS", "LOW")
    check_token()
    #We do valNum for projectID here because we need it in the comparison
    project_id = request.form['projectID']
    valNum(project_id, 12)
    project_name = request.form['projectName']
    valAlphaNum(project_name, 12)

    with contextlib.closing(get_db()) as con:
        #Check if submitted projectID is owned by user
        projects = con.execute('SELECT p.projectID, p.groupID, m.groupID, m.userID from projects as p JOIN groupMembers as m ON m.groupID = p.groupID where m.userID=?',
    				   [session['userID']]).fetchall()
        for project in projects:
            if int(project_id) == int(project[0]):
                f = request.form
                for key in f.keys():
                    for value in f.getlist(key):
                        if key.startswith(CKELEM_PREFIX):
                            idstr = key[CKELEM_PREFIX_LEN:]
                            listidx = "listID" + idstr
                            # The questionlist table has a listID field which is just an incrementing index.
                            # This form value carries the list name to be placed into the listName field.
                            list_name = request.form[listidx]
                            valAlphaNum(list_name, 12)
                            
                            answeridx = "answer" + idstr
                            answer = request.form[answeridx]
                            valAlphaNum(answer, 12)
                            
                            questionidx = "questionID" + idstr
                            questionID = request.form[questionidx]
                            valNum(questionID, 12)
                            
                            vulnidx = "vulnID" + idstr
                            vulnID = request.form[vulnidx]
                            valNum(vulnID, 12)
                            
                            date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
                            with con as cur:
                                cur.execute('INSERT INTO questionlist (entryDate, answer, projectName, projectID, questionID, vulnID, listName, userID) VALUES (?, ?, ?, ?, ?, ?, ?, ?)',
                                        [date, answer, project_name, project_id, questionID, vulnID, 
                                            list_name, session['userID']])
    redirect_url = "/results-checklists"
    return redirect(redirect_url)


def populate_checklists(checklist_paths, kb_paths, lvl_name, 
        owasp_ids_lvl, owasp_items_lvl, owasp_items_lvl_ygb, owasp_kb_ids_lvl, owasp_content_lvl, owasp_content_desc_lvl):
    for path in checklist_paths:
        basepath = os.path.basename(path)
        owasp_path_elems = basepath.split("-")
        owasp_id = get_num(owasp_path_elems[0])
        owasp_checklist_name = owasp_path_elems[2]
        if owasp_checklist_name == "ASVS":
            owasp_checklist_name = "-".join(owasp_path_elems[2:5])
            parse_index = 6
        else:
            parse_index = 4
        if owasp_checklist_name == lvl_name:
            owasp_kb = owasp_path_elems[parse_index]
            owasp_ygb = owasp_path_elems[parse_index + 2]

            owasp_ids_lvl.append(owasp_id)
            owasp_items_lvl.append(owasp_checklist_name)
            owasp_kb_ids_lvl.append(owasp_kb)
            owasp_items_lvl_ygb.append(owasp_ygb)
            with open(path, 'r') as pathf:
                checklistmd = pathf.read()
            owasp_content_lvl.append(Markup(markdown.markdown(checklistmd)))

            descriptions = []

            for kbpath in kb_paths:
                kbbasepath = os.path.basename(kbpath)
                path_vuln = get_num(kbbasepath.split("-")[0])
                if int(owasp_kb) == int(path_vuln):
                    with open(kbpath, 'r') as kbpathf:
                        kbmd = kbpathf.read()
                    description = kbmd.split("**")
                    descriptions.append(description[2])

            owasp_content_desc_lvl.append("\n".join(descriptions))
    return lvl_name
    

def populate_audit(checklist_paths, audit_name, audit_items, audit_ids, audit_kb_ids, audit_content):
    for path in checklist_paths:
        basepath = os.path.basename(path)
        audit_path_elems = basepath.split("-")
        audit_checklist_name = audit_path_elems[2]
        if audit_checklist_name == audit_name:
            audit_id = get_num(audit_path_elems[0])
            audit_kb = audit_path_elems[4]

            audit_ids.append(audit_id)
            audit_items.append(audit_checklist_name)
            audit_kb_ids.append(audit_kb)
            with open(path, 'r') as pathf:
                auditmd = pathf.read()
            audit_content.append(Markup(markdown.markdown(auditmd)))
    return audit_name


@app.route('/project-checklists/<project_id>', methods=['GET'])
@security
def project_checklists(project_id):
    """show the project checklists page"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /project-checklists", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    valNum(project_id, 12)
    with contextlib.closing(get_db()) as con:
        projects = con.execute('SELECT p.projectID, p.userID, p.groupID, p.projectName, p.projectVersion, p.projectDesc, p.ownerID, m.userID, m.groupID FROM projects as p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE p.projectID=? AND m.userID=?',
                            [project_id, session['userID']]).fetchall()
    projectName = ""
    owasp_items_lvl1 = []
    owasp_items_lvl1_ygb = []
    owasp_ids_lvl1 = []
    owasp_kb_ids_lvl1 = []
    owasp_content_lvl1 = []
    owasp_content_desc_lvl1 = []
    owasp_items_lvl2 = []
    owasp_items_lvl2_ygb = []
    owasp_ids_lvl2 = []
    owasp_kb_ids_lvl2 = []
    owasp_content_lvl2 = []
    owasp_content_desc_lvl2 = []
    owasp_items_lvl3 = []
    owasp_items_lvl3_ygb = []
    owasp_ids_lvl3 = []
    owasp_kb_ids_lvl3 = []
    owasp_content_lvl3 = []
    owasp_content_desc_lvl3 = []
    custom_items = []
    custom_ids = []
    custom_kb_ids = []
    custom_content = []
    basic_items = []
    basic_ids = []
    basic_kb_ids = []
    basic_content = []
    advanced_items = []
    advanced_ids = []
    advanced_kb_ids = []
    advanced_content = []

    for prep in projects:
        projectName = prep[3]
        
        checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
        checklist_paths.sort()

        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        kb_paths.sort()

        owasp_list_lvl1 = populate_checklists(checklist_paths, kb_paths, "ASVS-level-1",
                owasp_ids_lvl1, owasp_items_lvl1, owasp_items_lvl1_ygb, owasp_kb_ids_lvl1, owasp_content_lvl1, owasp_content_desc_lvl1)
        
        owasp_list_lvl2 = populate_checklists(checklist_paths, kb_paths, "ASVS-level-2",
                owasp_ids_lvl2, owasp_items_lvl2, owasp_items_lvl2_ygb, owasp_kb_ids_lvl2, owasp_content_lvl2, owasp_content_desc_lvl2)
        
        owasp_list_lvl3 = populate_checklists(checklist_paths, kb_paths, "ASVS-level-3",
                owasp_ids_lvl3, owasp_items_lvl3, owasp_items_lvl3_ygb, owasp_kb_ids_lvl3, owasp_content_lvl3, owasp_content_desc_lvl3)
        
        
        basic_list = populate_audit(checklist_paths, "CS_basic_audit",
                basic_items, basic_ids, basic_kb_ids, basic_content)
        
        advanced_list = populate_audit(checklist_paths, "CS_advanced_audit",
                advanced_items, advanced_ids, advanced_kb_ids, advanced_content)
        
        custom_list = populate_audit(checklist_paths, "custom",
                custom_items, custom_ids, custom_kb_ids, custom_content)

        break

    return render_template('project-checklists.html', csrf_token=session['csrf_token'], **locals())


@app.route('/results-checklists', methods=['GET'])
@security
def results_checklists():
    """show the results checklists page"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-checklists", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT q.answer, q.projectID, q.questionID,  q.vulnID, q.listName, q.entryDate, p.projectName, p.projectVersion, p.projectDesc, p.groupID, m.groupID, m.userID FROM questionlist AS q JOIN projects AS p ON q.projectID = p.projectID JOIN groupMembers as m ON m.groupID = p.groupID WHERE m.userID=? GROUP BY q.listName, q.entryDate ORDER BY p.projectName ASC',
                          [session['userID']]).fetchall()
    return render_template('results-checklists.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/results-functions', methods=['GET'])
@security
def results_functions():
    """show the results functions page"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-functions", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    with contextlib.closing(get_db()) as con:
        entries = con.execute('SELECT p.projectName, p.projectID, par.entryDate, p.projectDesc, p.groupID, m.userID, m.groupID, p.projectVersion, par.paramID, par.functionName, par.projectID FROM projects AS p join parameters AS par on p.projectID = par.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=? GROUP BY p.projectVersion ',
                         [session['userID']]).fetchall()
    return render_template('results-functions.html', entries=entries, csrf_token=session['csrf_token'])


@app.route('/results-functions-del', methods=['POST'])
@security
def functions_del():
    """delete functions result items"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-functions-del", "FAIL", "HIGH")
        abort(401)
    permissions("delete")
    check_token()
    entryDate = request.form['entryDate']
    projectID = request.form['projectID']
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        #Use select in order to see if this user is linked to project
        entries = con.execute("SELECT p.projectID, p.groupID, m.groupID, m.userID FROM projects AS p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=?  ",
                            [session['userID']]).fetchall()
        for entry in entries:
            if int(entry[0]) == int(projectID):
                with con as cur:
                    cur.execute("DELETE FROM parameters WHERE entryDate=? AND projectID=?",
                        [entryDate, projectID])
    return redirect("/results-functions")


@app.route('/results-checklists-del', methods=['POST'])
@security
def checklists_del():
    """delete checklist result item"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-checklists-del", "FAIL", "HIGH")
        abort(401)
    permissions("delete")
    check_token()
    entryDate = request.form['entryDate']
    projectID = request.form['projectID']
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        #Use select in order to see if this user is linked to project
        entries = con.execute("SELECT p.projectID, p.groupID, m.groupID, m.userID FROM projects AS p JOIN groupMembers AS m ON m.groupID = p.groupID WHERE m.userID=?  ",
                            [session['userID']]).fetchall()
        for entry in entries:
            if int(entry[0]) == int(projectID):
                with con as cur:
                    cur.execute("DELETE FROM questionlist WHERE entryDate=? AND projectID=? ",
                           [entryDate, projectID])
    return redirect("/results-checklists")


@app.route('/results-checklist-report/<entryDate>', methods=['GET'])
@security
def checklist_results(entryDate):
    """show checklist results report"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-checklist-report", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    ygb = []
    id_items = []
    questions = []
    content = []
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT l.listID, l.answer, l.projectID, l.projectName, l.questionID, l.vulnID, l.listName, l.entryDate, l.userID, m.userID, m.groupID, p.projectID, p.groupID FROM questionlist AS l JOIN projects AS p ON p.projectID = l.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE l.answer='no' AND l.entryDate=? AND m.userID=?",
               [entryDate, session['userID']]).fetchall()
    for entry in entries:
        projectName = entry[3]
        questionID = entry[4]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content.append(Markup(markdown.markdown(kbmd)))

                checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
                ckqs = []
                ckygbs = []
                for path in checklist_paths:
                    basepath = os.path.basename(path)
                    elems = basepath.split("-")
                    path_questionID = get_num(elems[0])
                    if int(questionID) == int(path_questionID):
                        with open(path, 'r') as pathf:
                            checklistmd = pathf.read()
                        questions.append(Markup(markdown.markdown(checklistmd)))
                        checklist_name = "-".join(elems[2:5])
                        if "ASVS" in checklist_name:
                            checklist_kb = elems[6]
                            checklist_ygb = elems[8]
                        else:
                            checklist_kb = elems[4]
                            checklist_ygb = elems[6]
                        ygb.append(checklist_ygb)
                questions.append("\n".join(ckqs))
                ygb.append("".join(ckygbs))

    return render_template('results-checklist-report.html', **locals())


@app.route('/results-checklist-docx/<entryDate>')
def download_file_checklist(entryDate):
    """Download checklist results report in docx"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-checklist-docx", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    ygb_docx = []
    content_raw = []
    content_checklist = []
    content_title = []
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT l.listID, l.answer, l.projectID, l.projectName, l.questionID, l.vulnID, l.listName, l.entryDate, l.userID, m.userID, m.groupID, p.projectID, p.groupID FROM questionlist AS l JOIN projects AS p ON p.projectID = l.projectID JOIN groupMembers AS m ON m.groupID = p.groupID WHERE l.answer='no' AND l.entryDate=? AND m.userID=?",
               [entryDate, session['userID']]).fetchall()
    document = Document()
    document.add_picture(os.path.join(app.root_path,'static/img/banner-docx.jpg'), width=Inches(5.125), height=Inches(1.042))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = entries[0][3]
    listName = entries[0][6]
    p.add_run('Used Checklist: '+listName)
    p.add_run('\r\n')
    p.add_run('Date: '+datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: '+projectName)
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    for entry in entries:
        projectName = entry[3]
        questionID = entry[4]
        vulnID = entry[5]
        listName = entry[6]
        entryDate = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content = Markup(markdown.markdown(kbmd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)

                checklist_paths = get_filepaths(os.path.join(app.root_path, "markdown/checklists"))
                kb_checklist_contents = []
                kb_checklist_ygb = []
                for path in checklist_paths:
                    basepath = os.path.basename(path)
                    elems = basepath.split("-")
                    path_questionID = get_num(elems[0])
                    if int(questionID) == int(path_questionID):
                        with open(path, 'r') as pathf:
                            checklistmd = pathf.read()
                        kb_checklist_contents.append(Markup(markdown.markdown(checklistmd)))
                        checklist_name = "-".join(elems[2:5])
                        if "ASVS" in checklist_name:
                            checklist_kb = elems[6]
                            checklist_ygb = elems[8]
                        else:
                            checklist_kb = elems[4]
                            checklist_ygb = elems[6]
                        kb_checklist_ygb.append(checklist_ygb)
                content_checklist.append("\n".join(kb_checklist_contents))
                ygb_docx.append("".join(kb_checklist_ygb))

    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrity of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In the post-development stage of the security knowledge framework the developer double-checks his application against a checklist which consists out of several questions asking the developer about different stages of development and the methodology of implementing different types of functionality the application contains. After filling in this checklist the developer gains feedback on the failed checklist items providing him with solutions about how to solve the additional vulnerability\'s found in the application.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading(content_title[i], level=1)
        result = re.sub("<p>", " ", content_checklist[i])
        result1 = re.sub("</p>", " ", result)
        document.add_heading(result1, level=4)
        p = document.add_paragraph(item.partition("\n")[2])
        for c in ygb_docx[i]:
            if c == "b":
                document.add_picture(os.path.join(app.root_path,'static/img/blue.png'), width=Inches(0.20))
            elif c == "g":
                document.add_picture(os.path.join(app.root_path,'static/img/green.png'), width=Inches(0.20))
            elif c == "y":
                document.add_picture(os.path.join(app.root_path,'static/img/yellow.png'), width=Inches(0.20))
        p.add_run("\n")
        document.add_page_break()
        i += 1
    # FIXME: save in a request-specific location or in memory
    document.save("checklist-security-report.docx")
    headers = {"Content-Disposition": "attachment; filename=%s" % "checklist-security-report.docx"}
    file_path = os.path.join(app.root_path, "checklist-security-report.docx")
    with open("checklist-security-report.docx", 'rb') as f:
        body = f.read()
    return make_response((body, headers))


@app.route('/results-function-report/<projectID>', methods=['GET'])
@security
def function_results(projectID):
    """show checklist results report"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-function-report", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    content = []
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate, parameters.techVuln, techhacks.techName, projects.userID, projects.groupID, m.userID, m.groupID FROM projects JOIN parameters ON parameters.projectID=projects.projectID JOIN techhacks ON techhacks.techID  = parameters.tech JOIN groupMembers AS m ON m.groupID = projects.groupID WHERE parameters.projectID=? AND m.userID=? GROUP BY parameters.tech;",
               [projectID, session['userID']]).fetchall()
    for entry in entries:
        projectName = entry[0]
        vulnID = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content.append(Markup(markdown.markdown(kbmd)))
    return render_template('results-function-report.html', **locals())


@app.route('/results-function-docx/<projectID>')
def download_file_function(projectID):
    """Download checklist results report in docx"""
    if not session.get('logged_in'):
        log( "User with no valid session tries access to page /results-function-docx", "FAIL", "HIGH")
        abort(401)
    permissions("read")
    content_raw = []
    content_title = []
    content_tech = []
    valNum(projectID, 12)
    with contextlib.closing(get_db()) as con:
        entries = con.execute("SELECT projects.projectName, projects.projectID, projects.projectVersion, parameters.functionName, parameters.tech, parameters.functionDesc, parameters.entryDate, parameters.techVuln, techhacks.techName, projects.userID, projects.groupID, m.userID, m.groupID FROM projects JOIN parameters ON parameters.projectID=projects.projectID JOIN techhacks ON techhacks.techID  = parameters.tech JOIN groupMembers AS m ON m.groupID = projects.groupID WHERE parameters.projectID=? AND m.userID=? GROUP BY parameters.tech;",
               [projectID, session['userID']]).fetchall()
    document = Document()
    document.add_picture(os.path.join(app.root_path,'static/img/banner-docx.jpg'), width=Inches(5.125), height=Inches(1.042))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #document.add_heading('Security Knowledge Framework', 0)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    projectName = entries[0][0]
    functionName = entries[0][3]
    functionDesc= entries[0][5]
    p.add_run('Date: '+datetime.datetime.now().strftime("%Y-%m-%d %H:%M"))
    p.add_run('\r\n')
    p.add_run('Project: '+projectName)
    document.add_page_break()
    p = document.add_heading('Table of contents', level=1)
    p.add_run('\r\n')
    document.add_paragraph('Introduction')
    for entry in entries:
        # entryDate = entry[6]
        vulnID = entry[7]
        kb_paths = get_filepaths(os.path.join(app.root_path, "markdown/knowledge_base"))
        for kbpath in kb_paths:
            kbbasepath = os.path.basename(kbpath)
            kbpath_vuln = get_num(kbbasepath.split("-")[0])
            if int(vulnID) == int(kbpath_vuln):
                with open(kbpath, 'r') as kbpathf:
                    kbmd = kbpathf.read()
                content = Markup(markdown.markdown(kbmd))
                text = ''.join(BeautifulSoup(content).findAll(text=True))
                text_encode = text.encode('utf-8')
                content_title.append(text_encode.splitlines()[0])
                text_encode = text_encode.replace("Solution", "\nSolution");
                content_raw.append(text_encode)
    for item in content_title:
        p = document.add_paragraph(item)
        p.add_run()
    document.add_page_break()
    document.add_heading('Introduction', level=1)
    p = document.add_paragraph(
        'The security knowledge framework is composed by means of the highest security standards currently available and is designed to maintain the integrity of your application, so you and your costumers sensitive data is protected against hackers. This document is provided with a checklist in which the programmers of your application had to run through in order to provide a secure product.'
    )
    p.add_run('\n')
    p = document.add_paragraph(
        'In this part of security knowledge framework, al the parameters and variables are audited by means of the information given by the programmer such as the processing techniques. Each of these techniques contain different types of vulnerabilities when implemented in a improper fashion. This document will raise awareness about these vulnerabilities, as well as presenting solutions for the right implementation.'
    )
    document.add_page_break()
    i = 0
    for item in content_raw:
        document.add_heading("Knowledge-Base: "+content_title[i], level=1)
        document.add_heading("Technology: "+entries[i][8], level=2)
        p = document.add_paragraph(item.partition("\n")[2])
        p.add_run("\n")
        document.add_page_break()
        i += 1
    # FIXME: save in a request-specific location such as memory
    document.save('function-security-report.docx')
    headers = {"Content-Disposition": "attachment; filename=%s" % "function-security-report.docx"}
    with open("function-security-report.docx", 'rb') as f:
        body = f.read()
    return make_response((body, headers))


if __name__ == "__main__":
    #Command line options to enable debug and/or saas (bind to 0.0.0.0)
    cmdargs = str(sys.argv)
    total = len(sys.argv)
    rand.cleanup()
    csrf_token_raw = rand.bytes(128)
    csrf_token = base64.b64encode(csrf_token_raw)
    for i in xrange(total):
        if (str(sys.argv[i][2:]) == "debug"):
            # Load default config and override config from an environment variable
            app.config.update(dict(
            DEBUG=True
            ))
        if (str(sys.argv[i][2:]) == "saas"):
            bindaddr = '0.0.0.0'
    if not os.path.isfile('server.crt'):
       app.run(host=bindaddr, port=5443, ssl_context='adhoc')
    else:
       context = SSL.Context(SSL.TLSv1_METHOD)
       context = ('server.crt', 'server.key')
       app.run(host=bindaddr, port=5443, ssl_context=context)
